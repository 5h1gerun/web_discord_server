from pathlib import Path
import os
import types
import sys
try:
    from PIL import Image
except Exception:
    Image = None
import pytest
sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

dummy_google = types.ModuleType("google.generativeai")
class _DummyCfg:
    def __init__(self, *a, **kw):
        pass

dummy_google.GenerationConfig = _DummyCfg
dummy_google.GenerativeModel = lambda *a, **kw: None
dummy_google.configure = lambda **kw: None
sys.modules.setdefault("google", types.ModuleType("google"))
sys.modules["google.generativeai"] = dummy_google

dummy_pdf2 = types.ModuleType("PyPDF2")

class _DummyPage:
    def extract_text(self):
        return ""

class _DummyPdf:
    def __init__(self, *a, **kw):
        self.pages = [_DummyPage()]

dummy_pdf2.PdfReader = _DummyPdf
sys.modules["PyPDF2"] = dummy_pdf2

if Image is None:
    dummy_pil = types.ModuleType("PIL")
    class _DummyImg:
        def save(self, *a, **k):
            import pathlib
            if a:
                p = pathlib.Path(a[0])
                p.parent.mkdir(parents=True, exist_ok=True)
                p.write_text("dummy")
    dummy_pil.Image = types.SimpleNamespace(open=lambda *a, **k: _DummyImg(), new=lambda *a, **k: _DummyImg())
    dummy_pil.UnidentifiedImageError = Exception
    sys.modules["PIL"] = dummy_pil
    sys.modules["PIL.Image"] = dummy_pil.Image

    # Fallback modules for optional dependencies
    dummy_docx = types.ModuleType("docx")
    class _DummyDoc:
        def __init__(self, *a, **k):
            self.paragraphs = []
        def add_paragraph(self, text=""):
            self.paragraphs.append(types.SimpleNamespace(text=text))
        def save(self, *a, **k):
            import pathlib
            if a:
                p = pathlib.Path(a[0])
                p.parent.mkdir(parents=True, exist_ok=True)
                p.write_text("dummy")
    dummy_docx.Document = _DummyDoc
    sys.modules.setdefault("docx", dummy_docx)
    sys.modules.setdefault("pptx", types.ModuleType("pptx")).Presentation = lambda *a, **k: types.SimpleNamespace(slides=[])
    sys.modules.setdefault("openpyxl", types.ModuleType("openpyxl")).load_workbook = lambda *a, **k: types.SimpleNamespace(worksheets=[])

import bot.auto_tag as auto_tag

class DummyResp:
    def __init__(self, text):
        self.text = text

class DummyModel:
    def __init__(self, *args, **kwargs):
        pass
    def generate_content(self, *args, **kwargs):
        return DummyResp("tagA, tagB")

def test_generate_tags(monkeypatch, tmp_path):
    f = tmp_path / "test.txt"
    f.write_text("sample text")
    dummy_genai = types.SimpleNamespace(
        configure=lambda **kw: None,
        GenerativeModel=lambda *a, **kw: DummyModel(),
    )
    monkeypatch.setenv("GEMINI_API_KEY", "dummy")
    monkeypatch.setattr(auto_tag, "genai", dummy_genai)
    tags = auto_tag.generate_tags(f)
    assert tags == "tagA, tagB"


def test_skip_large_file(monkeypatch, tmp_path):
    f = tmp_path / "big.txt"
    f.write_text("x")
    dummy_genai = types.SimpleNamespace(
        configure=lambda **kw: None,
        GenerativeModel=lambda *a, **kw: DummyModel(),
    )
    monkeypatch.setenv("GEMINI_API_KEY", "dummy")
    monkeypatch.setattr(auto_tag, "genai", dummy_genai)

    orig_stat = Path.stat

    def fake_stat(self):
        st = list(orig_stat(self))
        st[6] = 500_000_001
        return os.stat_result(st)

    monkeypatch.setattr(Path, "stat", fake_stat)
    tags = auto_tag.generate_tags(f)
    assert tags == ""


def test_skip_corrupted_file(monkeypatch, tmp_path):
    f = tmp_path / "bad.txt"
    f.write_text("x")
    dummy_genai = types.SimpleNamespace(
        configure=lambda **kw: None,
        GenerativeModel=lambda *a, **kw: DummyModel(),
    )
    monkeypatch.setenv("GEMINI_API_KEY", "dummy")
    monkeypatch.setattr(auto_tag, "genai", dummy_genai)

    def raise_error(*a, **k):
        raise OSError("broken")

    monkeypatch.setattr(Path, "read_bytes", raise_error)
    tags = auto_tag.generate_tags(f)
    assert tags == ""


@pytest.mark.skipif(Image is None, reason="Pillow not installed")
def test_generate_tags_binary(monkeypatch, tmp_path):
    f = tmp_path / "unknown.bin"
    img = Image.new("RGB", (1, 1), color="red")
    img.save(f, format="PNG")
    dummy_genai = types.SimpleNamespace(
        configure=lambda **kw: None,
        GenerativeModel=lambda *a, **kw: DummyModel(),
    )
    monkeypatch.setenv("GEMINI_API_KEY", "dummy")
    monkeypatch.setattr(auto_tag, "genai", dummy_genai)
    tags = auto_tag.generate_tags(f)
    assert tags == "tagA, tagB"


def test_generate_tags_docx(monkeypatch, tmp_path):
    try:
        from docx import Document
    except Exception:
        pytest.skip("python-docx not installed")
    doc = Document()
    doc.add_paragraph("hello docx")
    f = tmp_path / "test.docx"
    doc.save(f)
    dummy_genai = types.SimpleNamespace(
        configure=lambda **kw: None,
        GenerativeModel=lambda *a, **kw: DummyModel(),
    )
    monkeypatch.setenv("GEMINI_API_KEY", "dummy")
    monkeypatch.setattr(auto_tag, "genai", dummy_genai)
    tags = auto_tag.generate_tags(f)
    assert tags == "tagA, tagB"


def test_skip_zip_file(monkeypatch, tmp_path):
    import zipfile
    z = tmp_path / "test.zip"
    with zipfile.ZipFile(z, "w") as zp:
        zp.writestr("a.txt", "hello")
    dummy_genai = types.SimpleNamespace(
        configure=lambda **kw: None,
        GenerativeModel=lambda *a, **kw: DummyModel(),
    )
    monkeypatch.setenv("GEMINI_API_KEY", "dummy")
    monkeypatch.setattr(auto_tag, "genai", dummy_genai)
    tags = auto_tag.generate_tags(z)
    assert tags == ""


def test_skip_exe_file(monkeypatch, tmp_path):
    f = tmp_path / "app.exe"
    f.write_bytes(b"binary")
    dummy_genai = types.SimpleNamespace(
        configure=lambda **kw: None,
        GenerativeModel=lambda *a, **kw: DummyModel(),
    )
    monkeypatch.setenv("GEMINI_API_KEY", "dummy")
    monkeypatch.setattr(auto_tag, "genai", dummy_genai)
    tags = auto_tag.generate_tags(f)
    assert tags == ""
